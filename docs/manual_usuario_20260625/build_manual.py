from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "manual_usuario_20260625"
ASSET_DIR = OUT_DIR / "assets"
OUTPUT = OUT_DIR / "Manual_de_Usuario_Centro_IA_Video_20260625_actualizado.docx"

BRAND_GREEN = RGBColor(0x00, 0x68, 0x3A)
DARK_GREEN = RGBColor(0x00, 0x3D, 0x2A)
MUTED = RGBColor(0x5A, 0x6A, 0x64)
LIGHT_FILL = "F3F7F4"
HEADER_FILL = "E7F2EC"
BORDER = "CADBD2"


SCREENSHOTS = [
    (
        "00_login.png",
        "Pantalla de inicio de sesion",
        "Permite ingresar con correo corporativo y contrasena. Los usuarios deshabilitados no pueden acceder.",
    ),
    (
        "01_panel_principal.png",
        "Panel principal",
        "Resume el estado operativo de la plataforma y muestra accesos rapidos a los ultimos videos procesados.",
    ),
    (
        "02_gestion_videos.png",
        "Gestion de videos",
        "Concentra la carga de material, el historial reciente, la apertura de expedientes y las acciones administrativas permitidas.",
    ),
    (
        "03_modal_subir_video.png",
        "Carga de video y asignacion de area",
        "Solicita el archivo de video y la asignacion a un area/subarea antes de iniciar el procesamiento.",
    ),
    (
        "04_biblioteca.png",
        "Biblioteca por areas",
        "Muestra las areas disponibles en una columna y los videos del area seleccionada en el panel principal.",
    ),
    (
        "05_expediente_consulta.png",
        "Expediente del video: consulta",
        "Reune el reproductor, la consulta con fuentes y las acciones de mantenimiento autorizadas.",
    ),
    (
        "06_expediente_manuales.png",
        "Expediente del video: manuales",
        "Permite generar manuales con LLM local, revisar estado, previsualizar y descargar DOCX/PDF.",
    ),
    (
        "07_expediente_transcripcion.png",
        "Expediente del video: transcripcion",
        "Lista la transcripcion con timestamps; cada entrada puede llevar el reproductor al segundo correspondiente.",
    ),
    (
        "08_organizacion.png",
        "Organizacion de areas",
        "Administra la estructura de areas y subareas usada para clasificar videos y aplicar permisos.",
    ),
    (
        "09_usuarios.png",
        "Gestion de usuarios",
        "Administra cuentas, roles, estado de habilitacion e intentos fallidos de inicio de sesion.",
    ),
    (
        "10_modal_usuario.png",
        "Creacion de usuario",
        "Registra nombre, correo corporativo, rol asignado y contrasena temporal.",
    ),
    (
        "11_roles_permisos.png",
        "Roles y permisos",
        "Define permisos granulares y areas autorizadas; Super Admin es un rol del sistema.",
    ),
    (
        "12_modal_rol.png",
        "Creacion de rol",
        "Permite construir roles por combinacion de permisos y alcance de areas.",
    ),
]


PERMISSIONS = [
    ("Ver Dashboard", "Acceso al panel principal y resumen operativo."),
    ("Ver Gestion de Videos", "Acceso a la pagina de carga e historial de videos."),
    ("Ver Biblioteca", "Acceso al repositorio audiovisual organizado por areas."),
    ("Ver Organizacion", "Acceso a la estructura de areas y subareas."),
    ("Ver Usuarios", "Acceso a la pagina de usuarios."),
    ("Ver Roles", "Acceso a la pagina de roles y permisos."),
    ("Subir Videos", "Permite cargar nuevos videos y asignarlos durante la carga."),
    ("Editar Videos", "Permite cambiar nombre o asignacion de un video existente."),
    ("Reprocesar Videos", "Permite volver a extraer audio, transcribir e indexar un video."),
    ("Reindexar Videos", "Permite reconstruir el indice de busqueda sobre una transcripcion existente."),
    ("Eliminar Videos", "Permite eliminar un video y sus datos asociados."),
    ("Generar Manuales", "Permite crear manuales con LLM local y eliminar manuales existentes."),
    ("Gestionar Areas", "Permite crear areas y subareas."),
    ("Gestionar Usuarios", "Permite crear, editar, habilitar/deshabilitar y eliminar usuarios."),
    ("Gestionar Roles", "Permite crear, editar y eliminar roles no protegidos."),
]


VIDEO_STATUSES = [
    ("Subido", "El archivo fue recibido y queda pendiente de procesamiento."),
    ("Procesando", "El backend extrae audio, transcribe, detecta idioma e indexa contenido."),
    ("Listo", "El video tiene transcripcion e indice disponible para consulta."),
    ("Fallo", "El procesamiento no termino correctamente; revisar el error y reprocesar si corresponde."),
]


ACCENT_REPLACEMENTS = {
    "Indice": "Índice",
    "indice": "índice",
    "Guia": "Guía",
    "Guia ": "Guía ",
    "guia": "guía",
    "sesion": "sesión",
    "Sesion": "Sesión",
    "contrasena": "contraseña",
    "Contrasena": "Contraseña",
    "Gestion": "Gestión",
    "gestion": "gestión",
    "Administracion": "Administración",
    "administracion": "administración",
    "Organizacion": "Organización",
    "organizacion": "organización",
    "informacion": "información",
    "Informacion": "Información",
    "descripcion": "descripción",
    "Descripcion": "Descripción",
    "institucion": "institución",
    "Institucion": "Institución",
    "version": "versión",
    "Version": "Versión",
    "autenticacion": "autenticación",
    "Autenticacion": "Autenticación",
    "transcripcion": "transcripción",
    "Transcripcion": "Transcripción",
    "extraccion": "extracción",
    "Extraccion": "Extracción",
    "indexacion": "indexación",
    "Indexacion": "Indexación",
    "busqueda": "búsqueda",
    "Busqueda": "Búsqueda",
    "navegacion": "navegación",
    "Navegacion": "Navegación",
    "practicas": "prácticas",
    "Practicas": "Prácticas",
    "capacitacion": "capacitación",
    "Capacitacion": "Capacitación",
    "tecnico": "técnico",
    "Tecnico": "Técnico",
    "tecnica": "técnica",
    "Tecnica": "Técnica",
    "tecnicas": "técnicas",
    "Tecnicas": "Técnicas",
    "edicion": "edición",
    "Edicion": "Edición",
    "eliminacion": "eliminación",
    "Eliminacion": "Eliminación",
    "asincrono": "asíncrono",
    "Asincrono": "Asíncrono",
    "rapida": "rápida",
    "Rapida": "Rápida",
    "rapidas": "rápidas",
    "Rapidas": "Rápidas",
    "generacion": "generación",
    "Generacion": "Generación",
    "configuracion": "configuración",
    "Configuracion": "Configuración",
    "habilitacion": "habilitación",
    "Habilitacion": "Habilitación",
    "pagina": "página",
    "Pagina": "Página",
    "paginas": "páginas",
    "Paginas": "Páginas",
    "accion": "acción",
    "Accion": "Acción",
    "asignacion": "asignación",
    "Asignacion": "Asignación",
    "clasificacion": "clasificación",
    "Clasificacion": "Clasificación",
    "operacion": "operación",
    "Operacion": "Operación",
    "revision": "revisión",
    "Revision": "Revisión",
    "auditoria": "auditoría",
    "Auditoria": "Auditoría",
    "validacion": "validación",
    "Validacion": "Validación",
    "redaccion": "redacción",
    "Redaccion": "Redacción",
    "seccion": "sección",
    "Seccion": "Sección",
    "area": "área",
    "Area": "Área",
    "areas": "áreas",
    "Areas": "Áreas",
    "subarea": "subárea",
    "Subarea": "Subárea",
    "subareas": "subáreas",
    "Subareas": "Subáreas",
    "boton": "botón",
    "Boton": "Botón",
    "rapido": "rápido",
    "rapidos": "rápidos",
    "ultimo": "último",
    "Ultimo": "Último",
    "ultimos": "últimos",
    "Ultimos": "Últimos",
    "minimo": "mínimo",
    "Maximo": "Máximo",
    "maximo": "máximo",
    "segun": "según",
    "Segun": "Según",
    "tambien": "también",
    "Tambien": "También",
    "corresponda": "corresponda",
    "acciónes": "acciones",
    "Acciónes": "Acciones",
    "botónes": "botones",
    "Botónes": "Botones",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_picture_alt(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def set_run_font(run, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BRAND_GREEN, 18, 10),
        ("Heading 2", 13, BRAND_GREEN, 14, 7),
        ("Heading 3", 12, DARK_GREEN, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Centro IA Video - Manual de Usuario")
    set_run_font(run, size=9, color=MUTED)


def add_title_page(doc: Document) -> None:
    logo = ROOT / "frontend" / "src" / "assets" / "LogoBMSC.png"
    if logo.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_shape = p_logo.add_run().add_picture(str(logo), width=Inches(2.1))
        set_picture_alt(
            logo_shape,
            "Logo Banco Mercantil Santa Cruz",
            "Logotipo institucional utilizado en la portada del manual.",
        )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Manual de Usuario")
    set_run_font(r, size=26, color=DARK_GREEN, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("Centro IA Video")
    set_run_font(r, size=16, color=BRAND_GREEN, bold=True)

    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(28)
    r = description.add_run(
        "Guia operativa para administracion de videos, consultas, manuales, usuarios, roles y permisos."
    )
    set_run_font(r, size=11, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 7200, 120)
    mark_header_row(table.rows[0])
    rows = [
        ("Sistema", "Centro IA Video"),
        ("Institucion", "Banco Mercantil Santa Cruz"),
        ("Version del manual", "MVP actualizado - 25/06/2026"),
        ("Audiencia", "Gerencias, administradores y usuarios operativos autorizados"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, 2300 if idx == 0 else 4900)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()


def add_h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_p(doc: Document, text: str = "", bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 9360, 120)
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_border(cell, "9FC8B5")
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_GREEN, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, sum(widths), 120)
    mark_header_row(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_width(cell, widths[idx])
        set_cell_border(cell)
        set_cell_shading(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_GREEN)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = value
            set_cell_width(cell, widths[idx])
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, note: str) -> None:
    path = ASSET_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(6.2))
    set_picture_alt(shape, caption, note)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, color=DARK_GREEN, bold=True)
    detail = doc.add_paragraph()
    detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail.paragraph_format.space_after = Pt(12)
    r = detail.add_run(note)
    set_run_font(r, size=9, color=MUTED, italic=True)


def apply_spanish_accents(doc: Document) -> None:
    def update_paragraph(paragraph) -> None:
        for run in paragraph.runs:
            if not run.text:
                continue
            text = run.text
            for source, target in ACCENT_REPLACEMENTS.items():
                text = text.replace(source, target)
            run.text = text

    for paragraph in doc.paragraphs:
        update_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_paragraph(paragraph)


def add_toc(doc: Document) -> None:
    add_h(doc, "Indice", 1)
    items = [
        "1. Alcance y principios de uso",
        "2. Acceso y navegacion general",
        "3. Panel principal",
        "4. Gestion de videos",
        "5. Biblioteca audiovisual",
        "6. Expediente del video",
        "7. Organizacion de areas",
        "8. Usuarios, roles y permisos",
        "9. Seguridad de acceso",
        "10. Buenas practicas operativas",
    ]
    add_numbered(doc, items)
    doc.add_page_break()


def build_manual() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_toc(doc)

    add_h(doc, "1. Alcance y principios de uso", 1)
    add_p(
        doc,
        "Centro IA Video permite cargar videos de capacitacion, procesarlos localmente, consultar su contenido con evidencia temporal y generar manuales operativos a partir del material audiovisual.",
    )
    add_callout(
        doc,
        "Principio operativo",
        "Las respuestas, transcripciones y manuales se basan en el contenido procesado del video. El usuario debe validar la informacion final antes de publicarla como material institucional.",
    )
    add_bullets(
        doc,
        [
            "El procesamiento de videos usa extraccion de audio con FFmpeg, transcripcion con Faster-Whisper e indexacion local para busqueda.",
            "Las consultas al video recuperan fragmentos por similitud y muestran fuentes con timestamp para verificar la evidencia.",
            "La generacion de manuales usa LLM local por defecto y puede incluir capturas del video en puntos relevantes.",
            "El acceso a paginas, acciones y videos depende del rol del usuario y de las areas autorizadas.",
        ]
    )

    add_h(doc, "2. Acceso y navegacion general", 1)
    add_p(doc, "El ingreso se realiza con correo corporativo y contrasena asignada por un administrador.")
    add_figure(doc, *SCREENSHOTS[0])
    add_p(
        doc,
        "La barra lateral concentra las secciones principales: Panel principal, Gestion de videos, Biblioteca, Organizacion, Usuarios y Roles. Tambien incluye una biblioteca rapida para abrir expedientes recientes sin cambiar de seccion.",
    )
    add_bullets(
        doc,
        [
            "Las opciones visibles dependen de los permisos del rol.",
            "El chip 'Entorno institucional' indica que se esta operando dentro del sistema interno.",
            "El boton 'Cerrar sesion' termina la sesion del usuario actual.",
        ]
    )

    add_h(doc, "3. Panel principal", 1)
    add_p(
        doc,
        "El panel principal ofrece una vista ejecutiva del estado de la plataforma. Debe utilizarse para verificar rapidamente volumen de videos, material en procesamiento, fragmentos indexados y duracion total disponible.",
    )
    add_figure(doc, *SCREENSHOTS[1])
    add_table(
        doc,
        ["Elemento", "Uso"],
        [
            ("Total videos", "Cantidad de videos cargados y disponibles en el sistema."),
            ("En proceso", "Videos que todavia estan en cola, transcribiendo o indexando."),
            ("Transcripcion", "Segmentos y fragmentos indexados para consulta."),
            ("Actividad reciente", "Acceso rapido a los ultimos videos cargados."),
        ],
        [2500, 6860],
    )

    add_h(doc, "4. Gestion de videos", 1)
    add_p(
        doc,
        "La gestion de videos se usa para cargar nuevos materiales y revisar el historial reciente. Las acciones administrativas se muestran solamente si el rol posee el permiso correspondiente.",
    )
    add_figure(doc, *SCREENSHOTS[2])
    add_h(doc, "4.1 Carga de un video", 2)
    add_p(doc, "Para cargar un video se debe seleccionar el archivo y asignarlo al area/subarea correspondiente.")
    add_figure(doc, *SCREENSHOTS[3])
    add_numbered(
        doc,
        [
            "Ingresar a Gestion de videos.",
            "Seleccionar Subir video.",
            "Elegir el archivo. Se admiten videos MP4, MKV y otros formatos compatibles con FFmpeg.",
            "Seleccionar el area y subarea que corresponda al material.",
            "Confirmar la carga. El video quedara en procesamiento asincrono.",
        ],
    )
    add_table(doc, ["Estado", "Descripcion"], VIDEO_STATUSES, [2200, 7160])
    add_h(doc, "4.2 Acciones del historial", 2)
    add_table(
        doc,
        ["Accion", "Permiso requerido", "Descripcion"],
        [
            ("Subir video", "Subir Videos", "Permite registrar un nuevo archivo y asignarlo a un area."),
            ("Editar", "Editar Videos", "Permite modificar nombre visible o asignacion del video."),
            ("Eliminar", "Eliminar Videos", "Elimina el video y sus datos asociados. Debe usarse solo cuando corresponde retirar el material."),
            ("Abrir", "Ver Biblioteca o acceso al video", "Abre el expediente individual del video."),
        ],
        [1900, 2300, 5160],
    )

    add_h(doc, "5. Biblioteca audiovisual", 1)
    add_p(
        doc,
        "La biblioteca organiza los videos por areas. La columna izquierda muestra areas y subareas disponibles; el panel derecho muestra los videos del area seleccionada. Este diseno responde al modelo de permisos por area.",
    )
    add_figure(doc, *SCREENSHOTS[4])
    add_bullets(
        doc,
        [
            "El buscador filtra videos dentro del area seleccionada.",
            "Cada tarjeta muestra miniatura, estado, duracion, fecha, area asignada y cantidad de segmentos indexados.",
            "Abrir expediente lleva al detalle completo del video.",
            "Los usuarios con roles por area solo visualizan las areas autorizadas.",
        ]
    )

    add_h(doc, "6. Expediente del video", 1)
    add_p(
        doc,
        "El expediente es la pantalla central para revisar, consultar y documentar un video. Contiene reproductor, pestañas funcionales, resumen tecnico y acciones segun permisos.",
    )
    add_figure(doc, *SCREENSHOTS[5])
    add_table(
        doc,
        ["Zona", "Funcion"],
        [
            ("Reproductor", "Permite revisar el video y saltar al segundo exacto desde fuentes o transcripcion."),
            ("Consultar contenido", "Permite preguntar o buscar terminos dentro del video. La respuesta cita fragmentos como fuentes."),
            ("Manuales del Video", "Administra la generacion, vista previa y descarga de manuales."),
            ("Transcripcion", "Muestra segmentos con timestamp; cada segmento puede reproducirse desde su inicio."),
            ("Resumen tecnico", "Muestra estado, progreso, avance, motor de transcripcion, idioma, manuales y acciones administrativas."),
        ],
        [2500, 6860],
    )
    add_h(doc, "6.1 Consulta con fuentes", 2)
    add_p(
        doc,
        "El usuario escribe una pregunta o termino de busqueda. El sistema responde con una explicacion y muestra fuentes asociadas. Cada fuente conserva inicio y fin del fragmento para verificar la respuesta en el video.",
    )
    add_bullets(
        doc,
        [
            "Usar preguntas concretas mejora la precision.",
            "Cuando una fuente tiene boton de reproduccion, el reproductor salta al inicio del fragmento.",
            "La respuesta no reemplaza la revision del material cuando se prepara informacion oficial.",
        ]
    )
    add_h(doc, "6.2 Manuales y guias", 2)
    add_figure(doc, *SCREENSHOTS[6])
    add_p(
        doc,
        "La seccion de manuales genera documentos profesionales con LLM local. El usuario puede observar el estado, abrir la vista previa y descargar el resultado en DOCX o PDF.",
    )
    add_table(
        doc,
        ["Control", "Descripcion"],
        [
            ("LLM local", "Indica que la redaccion se genera con un modelo local configurado en el servidor."),
            ("Generar manual", "Inicia un nuevo manual del video listo. Requiere permiso Generar Manuales."),
            ("Ver", "Abre la previsualizacion renderizada del manual."),
            ("DOCX / PDF", "Descarga el manual para revision, distribucion o archivo."),
            ("Eliminar", "Retira un manual existente. Requiere permiso Generar Manuales."),
        ],
        [2500, 6860],
    )
    add_h(doc, "6.3 Transcripcion por timestamp", 2)
    add_figure(doc, *SCREENSHOTS[7])
    add_p(
        doc,
        "La transcripcion divide el contenido en segmentos con inicio exacto. Esta vista es util para auditoria, validacion de respuestas y revision puntual de contenido.",
    )
    add_h(doc, "6.4 Acciones tecnicas del expediente", 2)
    add_table(
        doc,
        ["Accion", "Permiso requerido", "Uso recomendado"],
        [
            ("Reindexar", "Reindexar Videos", "Usar cuando la transcripcion existe pero se requiere reconstruir la busqueda."),
            ("Reprocesar", "Reprocesar Videos", "Usar cuando se debe volver a extraer audio, transcribir e indexar."),
            ("Eliminar", "Eliminar Videos", "Usar solo para retirar material no valido o duplicado."),
        ],
        [1900, 2300, 5160],
    )

    add_h(doc, "7. Organizacion de areas", 1)
    add_p(
        doc,
        "Las areas y subareas clasifican el contenido y sirven como base para limitar el acceso de los roles. Antes de operar usuarios por area, la estructura debe estar creada correctamente.",
    )
    add_figure(doc, *SCREENSHOTS[8])
    add_bullets(
        doc,
        [
            "Nueva Area crea una unidad principal de clasificacion.",
            "Nueva subarea agrega divisiones dentro de un area.",
            "Las areas autorizadas se asignan desde la configuracion de roles.",
        ]
    )

    add_h(doc, "8. Usuarios, roles y permisos", 1)
    add_h(doc, "8.1 Gestion de usuarios", 2)
    add_p(
        doc,
        "La pagina de usuarios permite administrar cuentas corporativas, roles, estado de habilitacion e intentos fallidos. Las cuentas pueden editarse, eliminarse o deshabilitarse segun permisos.",
    )
    add_figure(doc, *SCREENSHOTS[9])
    add_figure(doc, *SCREENSHOTS[10])
    add_table(
        doc,
        ["Campo", "Regla operativa"],
        [
            ("Nombre completo", "Identifica al colaborador en la interfaz."),
            ("Correo corporativo", "Debe pertenecer al dominio institucional @bmsc.com.bo."),
            ("Rol asignado", "Determina paginas, acciones y areas visibles."),
            ("Contrasena temporal", "Se define al crear el usuario o al resetear acceso."),
            ("Habilitado/Deshabilitado", "Un usuario deshabilitado no puede iniciar sesion."),
        ],
        [2500, 6860],
    )
    add_callout(
        doc,
        "Regla de correo",
        "No puede existir el mismo correo en usuarios habilitados o deshabilitados. Si el usuario fue eliminado, el correo puede volver a registrarse.",
    )
    add_h(doc, "8.2 Roles y permisos", 2)
    add_p(
        doc,
        "Los roles agrupan permisos y definen el alcance por areas. Un rol puede tener acceso global a todas las areas o acceso restringido a areas especificas.",
    )
    add_figure(doc, *SCREENSHOTS[11])
    add_figure(doc, *SCREENSHOTS[12])
    add_p(
        doc,
        "El rol Super Admin es un rol del sistema: posee acceso global, todos los permisos, no muestra contador de intentos fallidos y no debe editarse ni eliminarse.",
    )
    add_table(doc, ["Permiso", "Alcance"], PERMISSIONS, [2700, 6660])
    add_callout(
        doc,
        "Eliminacion de roles",
        "Un rol no puede eliminarse si esta asignado a usuarios. El sistema debe informar que el rol esta en uso y listar los usuarios asociados para que el administrador reasigne o elimine esas cuentas antes de retirar el rol.",
    )

    add_h(doc, "9. Seguridad de acceso", 1)
    add_p(
        doc,
        "La seguridad combina autenticacion, estado de cuenta, permisos granulares y restricciones por area.",
    )
    add_table(
        doc,
        ["Mecanismo", "Comportamiento"],
        [
            ("Sesion autenticada", "Todas las paginas internas requieren usuario autenticado."),
            ("Permisos por pagina", "La navegacion oculta o bloquea paginas sin permiso."),
            ("Permisos por accion", "Editar, reindexar, reprocesar, eliminar y generar manuales requieren permisos especificos."),
            ("Areas autorizadas", "Los roles por area solo ven videos del area permitida."),
            ("Bloqueo por intentos", "Usuarios que no son Super Admin quedan deshabilitados al quinto intento fallido."),
            ("Rehabilitacion", "Solo un Super Admin puede habilitar nuevamente usuarios deshabilitados."),
        ],
        [2700, 6660],
    )
    add_bullets(
        doc,
        [
            "Super Admin no esta sujeto al contador de intentos fallidos.",
            "Un usuario no puede deshabilitarse a si mismo.",
            "Los usuarios deshabilitados no pueden iniciar sesion hasta ser habilitados por un Super Admin.",
            "La eliminacion de usuario es diferente a deshabilitar: al eliminar, el correo puede reutilizarse.",
        ]
    )

    add_h(doc, "10. Buenas practicas operativas", 1)
    add_bullets(
        doc,
        [
            "Asignar siempre el area y subarea correcta al subir un video para que la biblioteca y los permisos funcionen correctamente.",
            "Esperar estado Listo antes de consultar, revisar transcripcion o generar manuales.",
            "Usar Reindexar solo cuando la transcripcion ya existe y la busqueda necesita actualizarse.",
            "Usar Reprocesar cuando se sospecha un problema en audio, transcripcion o segmentacion.",
            "Validar los manuales generados antes de distribuirlos formalmente.",
            "Mantener roles con el minimo permiso necesario para cada perfil de usuario.",
            "Deshabilitar usuarios cuando se desea impedir acceso temporal; eliminar solo cuando la cuenta ya no debe conservarse.",
        ]
    )
    add_h(doc, "Anexo. Flujo operativo recomendado", 1)
    add_numbered(
        doc,
        [
            "Crear o validar areas y subareas.",
            "Crear roles con permisos y areas autorizadas.",
            "Crear usuarios y asignar rol.",
            "Cargar video y asignarlo al area correspondiente.",
            "Esperar finalizacion de procesamiento.",
            "Consultar contenido y validar fuentes.",
            "Generar manual si corresponde.",
            "Descargar y revisar DOCX/PDF antes de publicarlo.",
        ],
    )

    doc.core_properties.title = "Manual de Usuario - Centro IA Video"
    doc.core_properties.subject = "Guia operativa del sistema Centro IA Video"
    doc.core_properties.author = "Banco Mercantil Santa Cruz"
    doc.core_properties.created = datetime.now()
    apply_spanish_accents(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_manual()
