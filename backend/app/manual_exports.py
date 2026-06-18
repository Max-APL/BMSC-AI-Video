from __future__ import annotations

import html
import re
from pathlib import Path
from typing import List, Optional, Tuple


class ManualExportError(RuntimeError):
    pass


def export_manual(
    content: str,
    output_path: Path,
    output_format: str,
    *,
    assets_dir: Optional[Path] = None,
) -> Path:
    output_format = output_format.lower()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_format == "markdown":
        output_path.write_text(content, encoding="utf-8")
        return output_path
    if output_format == "docx":
        return export_docx(content, output_path, assets_dir=assets_dir)
    if output_format == "pdf":
        return export_pdf(content, output_path, assets_dir=assets_dir)
    raise ManualExportError("Formato de descarga no soportado. Usa markdown, docx o pdf.")


def export_docx(content: str, output_path: Path, *, assets_dir: Optional[Path] = None) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise ManualExportError(
            "No se pudo generar DOCX porque falta python-docx. "
            "Instala dependencias con: pip install -r requirements.txt"
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 4"].font.name = "Arial"
    styles["Heading 4"].font.size = Pt(11)

    for kind, text in parse_markdown_blocks(content):
        if kind == "h1":
            paragraph = document.add_heading(text, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "h2":
            document.add_heading(text, level=2)
        elif kind == "h3":
            document.add_heading(text, level=3)
        elif kind == "h4":
            document.add_heading(text, level=4)
        elif kind == "bullet":
            add_formatted_paragraph(document, text, style="List Bullet")
        elif kind == "number":
            add_formatted_paragraph(document, text, style="List Number")
        elif kind == "image":
            add_docx_image(document, text, assets_dir)
        elif kind == "paragraph":
            add_formatted_paragraph(document, text)

    document.save(output_path)
    return output_path


def export_pdf(content: str, output_path: Path, *, assets_dir: Optional[Path] = None) -> Path:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import HRFlowable, Image, Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise ManualExportError(
            "No se pudo generar PDF porque falta reportlab. "
            "Instala dependencias con: pip install -r requirements.txt"
        ) from exc

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ManualTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#12395f"),
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ManualMeta",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#60748c"),
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ManualHeading2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#12395f"),
            spaceBefore=12,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ManualHeading3",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=14,
            textColor=colors.HexColor("#1f6ed4"),
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ManualHeading4",
            parent=styles["Heading4"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=colors.HexColor("#263b52"),
            spaceBefore=6,
            spaceAfter=4,
        )
    )
    body_style = ParagraphStyle(
        name="ManualBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        spaceAfter=7,
    )
    bullet_style = ParagraphStyle(
        name="ManualBullet",
        parent=body_style,
        leftIndent=18,
        firstLineIndent=-10,
        spaceAfter=5,
    )
    number_style = ParagraphStyle(
        name="ManualNumber",
        parent=body_style,
        leftIndent=20,
        firstLineIndent=-14,
        spaceAfter=5,
    )
    caption_style = ParagraphStyle(
        name="ManualCaption",
        parent=body_style,
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#60748c"),
        alignment=1,
        spaceAfter=10,
    )

    story = []
    number_index = 1
    seen_title = False

    for kind, text in parse_markdown_blocks(content):
        if kind == "h1":
            number_index = 1
            story.append(Paragraph(inline_markdown(text), styles["ManualTitle"]))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dbe5f0"), spaceAfter=14))
            seen_title = True
        elif kind == "h2":
            number_index = 1
            if seen_title and text.lower() == "datos del documento":
                story.append(Paragraph(inline_markdown(text), styles["ManualHeading2"]))
            else:
                story.append(Spacer(1, 4))
                story.append(Paragraph(inline_markdown(text), styles["ManualHeading2"]))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e3ebf4"), spaceAfter=6))
            seen_title = True
        elif kind == "h3":
            number_index = 1
            story.append(Paragraph(inline_markdown(text), styles["ManualHeading3"]))
        elif kind == "h4":
            number_index = 1
            story.append(Paragraph(inline_markdown(text), styles["ManualHeading4"]))
        elif kind == "bullet":
            story.append(Paragraph("&bull;&nbsp;" + inline_markdown(text), bullet_style))
        elif kind == "number":
            story.append(Paragraph(f"{number_index}.&nbsp;" + inline_markdown(text), number_style))
            number_index += 1
        elif kind == "image":
            image_path, caption = resolve_image_payload(text, assets_dir)
            if image_path:
                image = Image(str(image_path))
                image._restrictSize(6.7 * inch, 3.75 * inch)
                story.append(image)
                if caption:
                    story.append(Paragraph(inline_markdown(caption), caption_style))
        elif kind == "paragraph":
            number_index = 1
            story.append(Paragraph(inline_markdown(text), body_style))

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Manual de capacitacion",
    )
    document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return output_path


def parse_markdown_blocks(content: str) -> List[Tuple[str, str]]:
    blocks: List[Tuple[str, str]] = []
    paragraph_lines: List[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            blocks.append(("paragraph", " ".join(paragraph_lines).strip()))
            paragraph_lines = []

    raw_lines = content.splitlines()
    index = 0
    while index < len(raw_lines):
        raw_line = raw_lines[index]
        line = raw_line.strip()
        next_line = raw_lines[index + 1].strip() if index + 1 < len(raw_lines) else ""
        if not line:
            flush_paragraph()
            index += 1
            continue
        if next_line and re.fullmatch(r"={3,}|-{3,}", next_line):
            flush_paragraph()
            blocks.append(("h2", line.strip("# *")))
            index += 2
            continue
        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if image_match:
            flush_paragraph()
            blocks.append(("image", f"{image_match.group(1).strip()}\n{image_match.group(2).strip()}"))
            index += 1
            continue
        if line.startswith("# "):
            flush_paragraph()
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            flush_paragraph()
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            flush_paragraph()
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("#### "):
            flush_paragraph()
            blocks.append(("h4", line[5:].strip()))
        elif line.startswith("- ") or line.startswith("* ") or line.startswith("+ "):
            flush_paragraph()
            blocks.append(("bullet", line[2:].strip()))
        elif re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            blocks.append(("number", re.sub(r"^\d+\.\s+", "", line).strip()))
        else:
            paragraph_lines.append(line)
        index += 1
    flush_paragraph()
    return blocks


def add_formatted_paragraph(document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_docx_image(document, payload: str, assets_dir: Optional[Path]) -> None:
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
    except ImportError as exc:
        raise ManualExportError(
            "No se pudo generar DOCX porque falta python-docx. "
            "Instala dependencias con: pip install -r requirements.txt"
        ) from exc

    image_path, caption = resolve_image_payload(payload, assets_dir)
    if not image_path:
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    if caption:
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run_item in caption_paragraph.runs:
            run_item.italic = True


def resolve_image_payload(payload: str, assets_dir: Optional[Path]) -> Tuple[Optional[Path], str]:
    lines = payload.splitlines()
    caption = lines[0].strip() if lines else ""
    raw_path = lines[1].strip() if len(lines) > 1 else ""
    if not raw_path or not assets_dir:
        return None, caption

    normalized = raw_path.replace("\\", "/").strip("/")
    parts = [part for part in normalized.split("/") if part and part not in {".", ".."}]
    if not parts or len(parts) != len([part for part in normalized.split("/") if part]):
        return None, caption

    root = assets_dir.resolve()
    image_path = (root / Path(*parts)).resolve()
    if root not in image_path.parents and image_path != root:
        return None, caption
    if not image_path.exists() or not image_path.is_file():
        return None, caption
    return image_path, caption


def inline_markdown(text: str) -> str:
    escaped = html.escape(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)


def draw_footer(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColorRGB(0.38, 0.45, 0.55)
    canvas.drawRightString(document.pagesize[0] - document.rightMargin, 0.45 * 72, f"Pagina {document.page}")
    canvas.restoreState()
